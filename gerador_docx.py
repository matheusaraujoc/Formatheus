# gerador_docx.py
# Descrição: Versão com a correção do recuo (TAB) na primeira
# linha do Resumo.
#
# ATUALIZAÇÃO (v47 - Gráficos 3D):
# 1. Adicionada lógica para renderizar {{Grafico3D:Titulo}}.
# 2. Importado Grafico3D e adicionado contador_graficos_3d.
# 3. Atualizado regex e adicionado método _renderizar_grafico_3d.
#

import os
import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import win32com.client as win32
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False
    print("AVISO: Biblioteca 'pywin32' não encontrada. A automação do sumário será desativada.")

# --- INÍCIO DA MODIFICAÇÃO (v47) ---
from documento import (DocumentoABNT, Capitulo, ItemLista, ListaABNT, 
                       Grafico, Grafico3D) # <--- ADICIONADO Grafico3D
# --- FIM DA MODIFICAÇÃO ---
from normas_abnt import MotorNormasABNT

def adicionar_sumario(doc, paragrafo_placeholder):
    sdt = OxmlElement('w:sdt')
    sdtContent = OxmlElement('w:sdtContent')
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_separate)
    r.append(fldChar_end)
    p.append(r)
    sdtContent.append(p)
    sdt.append(sdtContent)
    p_xml = paragrafo_placeholder._p
    p_xml.addnext(sdt)
    p_xml.getparent().remove(p_xml)

class GeradorDOCX:
    def __init__(self, doc_abnt: DocumentoABNT):
        self.doc_abnt = doc_abnt
        self.doc = Document()
        self.regras = MotorNormasABNT(self.doc_abnt)
        self.regras.configurar_pagina_e_estilos(self.doc)
        self.contador_tabelas = 0
        self.contador_figuras = 0
        self.contador_formulas = 0
        self.contador_listas = 0 
        self.contador_graficos = 0
        self.contador_graficos_3d = 0 # <--- NOVO CONTADOR (v47)

    def _atualizar_sumario_com_word(self, caminho_arquivo):
        if not WIN32_AVAILABLE:
            print("Não foi possível atualizar o sumário: pywin32 não está instalado.")
            return False
        word = None
        try:
            print("Iniciando automação do MS Word para reconstrução do sumário...")
            word = win32.DispatchEx("Word.Application")
            word.Visible = False
            doc_path = os.path.abspath(caminho_arquivo)
            doc = word.Documents.Open(doc_path)
            if doc.TablesOfContents.Count > 0:
                doc.TablesOfContents(1).Update()
            doc.Save()
            doc.Close(SaveChanges=False)
            print("Sumário clicável gerado e atualizado com sucesso.")
            return True
        except Exception as e:
            print(f"ERRO: Falha ao automatizar o Word para atualizar o sumário: {e}")
            return False
        finally:
            if word is not None:
                word.Quit()

    def gerar_documento(self, caminho_arquivo: str):
        if self.regras.is_artigo:
            self._gerar_artigo(caminho_arquivo)
        else:
            self._gerar_trabalho_academico(caminho_arquivo)

    def _gerar_trabalho_academico(self, caminho_arquivo: str):
        
        # Seção 1: Capa (é a self.doc.sections[0])
        self._renderizar_capa() 
        
        # Seção 2: Folha de Rosto
        section_rosto = self.doc.add_section(WD_SECTION.NEW_PAGE)
        # Desvincula o rodapé da Seção 2 (Folha de Rosto) da Seção 1 (Capa)
        section_rosto.footer.is_linked_to_previous = False
        self._renderizar_folha_rosto(section_rosto) # Passa a seção como argumento

        # Seção 3: Resumo
        section_resumo = self.doc.add_section(WD_SECTION.NEW_PAGE)
        # Desvincula o rodapé da Seção 3 (Resumo)
        section_resumo.footer.is_linked_to_previous = False
        self._renderizar_resumo()

        # Seção 4: Conteúdo principal (onde a numeração começa)
        section_main = self.doc.add_section(WD_SECTION.NEW_PAGE)
        # Desvincula o rodapé da Seção 4 (Conteúdo)
        section_main.footer.is_linked_to_previous = False
        
        # A numeração de página é aplicada APENAS na Seção 4
        self._set_page_numbering(section_main) 
        
        self._renderizar_sumario()
        self._renderizar_secoes_recursivamente(self.doc_abnt.estrutura_textual)
        
        # Seção 5: Referências
        section_refs = self.doc.add_section(WD_SECTION.NEW_PAGE)
        # Mantém o cabeçalho/rodapé da seção anterior (com número de página)
        section_refs.header.is_linked_to_previous = True
        section_refs.footer.is_linked_to_previous = True
        
        self._renderizar_referencias()

        self.doc.save(caminho_arquivo)
        self._atualizar_sumario_com_word(caminho_arquivo)

    def _gerar_artigo(self, caminho_arquivo: str):
        section = self.doc.sections[0]
        self._set_page_numbering(section)
        self.regras.renderizar_cabecalho_artigo(self.doc)
        self._renderizar_secoes_recursivamente(self.doc_abnt.estrutura_textual)
        self.doc.add_section(WD_SECTION.NEW_PAGE)
        self._renderizar_referencias()

        self.doc.save(caminho_arquivo)
        print("Documento de Artigo Científico gerado com sucesso.")

    def _set_no_border_to_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

    def _renderizar_secoes_recursivamente(self, no_pai: Capitulo, prefixo_numeracao=""):
        for i, no_filho in enumerate(no_pai.filhos, 1):
            numero_completo = f"{prefixo_numeracao}{i}"
            nivel_titulo = len(numero_completo.split('.'))
            
            if nivel_titulo == 1 and i > 1:
                self.doc.add_page_break()
            
            self.regras.aplicar_estilo_titulo_secao(self.doc, numero_completo, no_filho.titulo, nivel=nivel_titulo)

            if no_filho.conteudo:
                # --- REGEX ATUALIZADO (v47) ---
                padrao = r"\{\{(?:(Tabela|Figura|Formula|Lista|Grafico|Grafico3D):([^}]+)|(QUEBRA_PAGINA|PAGINA_EM_BRANCO))\}\}"
                # --- FIM DA ATUALIZAÇÃO ---
                partes = re.split(padrao, no_filho.conteudo)

                idx = 0
                while idx < len(partes):
                    bloco_de_texto = partes[idx]
                    if bloco_de_texto and bloco_de_texto.strip():
                        paragrafos = bloco_de_texto.strip().split('\n')
                        for texto_paragrafo in paragrafos:
                            if texto_paragrafo.strip():
                                p = self.doc.add_paragraph()
                                self.regras.aplicar_estilo_paragrafo_normal(p, texto_paragrafo)
                    
                    if idx + 3 < len(partes): 
                        tipo_obj = partes[idx+1]
                        titulo_obj = partes[idx+2]
                        comando = partes[idx+3]
                        
                        if tipo_obj and titulo_obj:
                            if tipo_obj == "Tabela":
                                obj = next((t for t in self.doc_abnt.banco_tabelas if t.titulo == titulo_obj), None)
                                if obj:
                                    self.contador_tabelas += 1
                                    obj.numero = self.contador_tabelas
                                    self._renderizar_tabela(obj)
                            elif tipo_obj == "Figura":
                                obj = next((f for f in self.doc_abnt.banco_figuras if f.titulo == titulo_obj), None)
                                if obj:
                                    self.contador_figuras += 1
                                    obj.numero = self.contador_figuras
                                    self._renderizar_figura(obj)
                            
                            elif tipo_obj == "Grafico":
                                obj = next((g for g in self.doc_abnt.banco_graficos if g.titulo == titulo_obj), None)
                                if obj:
                                    self.contador_graficos += 1
                                    obj.numero = self.contador_graficos
                                    self._renderizar_grafico(obj) 
                            
                            # --- INÍCIO: LÓGICA DE RENDERIZAÇÃO DO GRÁFICO 3D (v47) ---
                            elif tipo_obj == "Grafico3D":
                                obj = next((g for g in self.doc_abnt.banco_graficos_3d if g.titulo == titulo_obj), None)
                                if obj:
                                    self.contador_graficos_3d += 1
                                    obj.numero = self.contador_graficos_3d
                                    self._renderizar_grafico_3d(obj) # Chama o novo método
                            # --- FIM: LÓGICA DE RENDERIZAÇÃO DO GRÁFICO 3D (v47) ---

                            elif tipo_obj == "Formula":
                                obj = next((f for f in self.doc_abnt.banco_formulas if f.legenda == titulo_obj), None)
                                if obj:
                                    self.contador_formulas += 1
                                    obj.numero = self.contador_formulas
                                    self._renderizar_formula(obj)
                            
                            elif tipo_obj == "Lista":
                                obj = next((l for l in self.doc_abnt.banco_listas if l.titulo == titulo_obj), None)
                                if obj:
                                    self.contador_listas += 1
                                    obj.numero = self.contador_listas
                                    self._renderizar_lista(obj)
                            
                        elif comando:
                            if comando == "QUEBRA_PAGINA":
                                self.doc.add_page_break()
                            elif comando == "PAGINA_EM_BRANCO":
                                self.doc.add_page_break()
                                p = self.doc.add_paragraph()
                                run = p.add_run()
                                run.text = " " 
                                p.paragraph_format.space_after = Pt(0)
                                self.doc.add_page_break()
                                
                        idx += 4 # Avança (Texto, Tipo, Titulo, Comando)
                    else:
                        idx += 1 # Avança o último bloco de texto
                
            self._renderizar_secoes_recursivamente(no_filho, prefixo_numeracao=f"{numero_completo}.")

    def _renderizar_tabela(self, tabela_obj):
        p_titulo = self.doc.add_paragraph()
        p_titulo.add_run(f"Tabela {tabela_obj.numero} – {tabela_obj.titulo}")
        self.regras.aplicar_estilo_legenda(p_titulo, is_titulo=True)
        p_titulo.paragraph_format.keep_with_next = True

        if not tabela_obj.dados: return

        num_rows = len(tabela_obj.dados)
        num_cols = len(tabela_obj.dados[0]) if num_rows > 0 else 0
        t = self.doc.add_table(rows=num_rows, cols=num_cols)
        t.style = 'Table Grid'

        for i, row_data in enumerate(tabela_obj.dados):
            for j, cell_data in enumerate(row_data):
                cell = t.cell(i, j)
                cell.text = cell_data
                p = cell.paragraphs[0]
                run = p.runs[0]
                run.font.name = self.regras.FONTE_PADRAO
                run.font.size = self.regras.TAMANHO_FONTE_LEGENDA
                run.font.color.rgb = self.regras.COR_FONTE_PADRAO
                
                if i == 0:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                else:
                    if tabela_obj.centralizar_conteudo:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    else:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        if tabela_obj.estilo_borda == 'abnt':
            self.regras.aplicar_estilo_tabela_abnt(t)

        if tabela_obj.fonte:
            p_fonte = self.doc.add_paragraph()
            p_fonte.add_run(f"Fonte: {tabela_obj.fonte}")
            self.regras.aplicar_estilo_legenda(p_fonte, is_titulo=False)

    def _renderizar_figura(self, figura_obj):
        p_titulo = self.doc.add_paragraph()
        p_titulo.add_run(f"Figura {figura_obj.numero} – {figura_obj.titulo}")
        self.regras.aplicar_estilo_legenda(p_titulo, is_titulo=True)
        p_titulo.paragraph_format.keep_with_next = True

        p_imagem = self.doc.add_paragraph()
        p_imagem.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        try:
            p_imagem.add_run().add_picture(figura_obj.caminho_processado, width=Cm(figura_obj.largura_cm))
        except Exception as e:
            run_erro = p_imagem.add_run(f"[ERRO: Imagem '{figura_obj.caminho_processado}' não encontrada ou inválida. {e}]")
            run_erro.italic = True
            run_erro.font.color.rgb = self.regras.COR_FONTE_PADRAO

        p_imagem.paragraph_format.space_before = Pt(0)
        p_imagem.paragraph_format.space_after = Pt(0)
        p_imagem.paragraph_format.keep_with_next = True

        if figura_obj.fonte:
            p_fonte = self.doc.add_paragraph()
            p_fonte.add_run(f"Fonte: {figura_obj.fonte}")
            self.regras.aplicar_estilo_legenda(p_fonte, is_titulo=False)
        else:
            p_imagem.paragraph_format.keep_with_next = False

    def _renderizar_grafico(self, grafico_obj):
        """Renderiza um gráfico (imagem) no DOCX, formatado como uma Figura."""
        
        # 1. Título (Legenda Superior)
        p_titulo = self.doc.add_paragraph()
        p_titulo.add_run(f"Gráfico {grafico_obj.numero} – {grafico_obj.titulo}")
        self.regras.aplicar_estilo_legenda(p_titulo, is_titulo=True)
        p_titulo.paragraph_format.keep_with_next = True

        # 2. Imagem (Centralizada)
        p_imagem = self.doc.add_paragraph()
        p_imagem.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        try:
            p_imagem.add_run().add_picture(grafico_obj.caminho_imagem_processada, width=Cm(grafico_obj.largura_cm))
        except Exception as e:
            run_erro = p_imagem.add_run(f"[ERRO: Imagem do Gráfico '{grafico_obj.caminho_imagem_processada}' não encontrada. {e}]")
            run_erro.italic = True
            run_erro.font.color.rgb = self.regras.COR_FONTE_PADRAO

        p_imagem.paragraph_format.space_before = Pt(0)
        p_imagem.paragraph_format.space_after = Pt(0)
        p_imagem.paragraph_format.keep_with_next = True

        # 3. Fonte (Legenda Inferior)
        if grafico_obj.fonte:
            p_fonte = self.doc.add_paragraph()
            p_fonte.add_run(f"Fonte: {grafico_obj.fonte}")
            self.regras.aplicar_estilo_legenda(p_fonte, is_titulo=False)
        else:
            p_imagem.paragraph_format.keep_with_next = False

    # --- INÍCIO: NOVO MÉTODO PARA RENDERIZAR GRÁFICO 3D (v47) ---
    def _renderizar_grafico_3d(self, grafico_obj):
        """Renderiza um gráfico 3D (imagem) no DOCX, formatado como uma Figura."""
        
        # 1. Título (Legenda Superior)
        p_titulo = self.doc.add_paragraph()
        p_titulo.add_run(f"Gráfico 3D {grafico_obj.numero} – {grafico_obj.titulo}")
        self.regras.aplicar_estilo_legenda(p_titulo, is_titulo=True)
        p_titulo.paragraph_format.keep_with_next = True

        # 2. Imagem (Centralizada)
        p_imagem = self.doc.add_paragraph()
        p_imagem.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        try:
            p_imagem.add_run().add_picture(grafico_obj.caminho_imagem_processada, width=Cm(grafico_obj.largura_cm))
        except Exception as e:
            run_erro = p_imagem.add_run(f"[ERRO: Imagem do Gráfico 3D '{grafico_obj.caminho_imagem_processada}' não encontrada. {e}]")
            run_erro.italic = True
            run_erro.font.color.rgb = self.regras.COR_FONTE_PADRAO

        p_imagem.paragraph_format.space_before = Pt(0)
        p_imagem.paragraph_format.space_after = Pt(0)
        p_imagem.paragraph_format.keep_with_next = True

        # 3. Fonte (Legenda Inferior)
        if grafico_obj.fonte:
            p_fonte = self.doc.add_paragraph()
            p_fonte.add_run(f"Fonte: {grafico_obj.fonte}")
            self.regras.aplicar_estilo_legenda(p_fonte, is_titulo=False)
        else:
            p_imagem.paragraph_format.keep_with_next = False
    # --- FIM: NOVO MÉTODO (v47) ---

    def _renderizar_formula(self, formula_obj):
        p_formula = self.doc.add_paragraph()

        tab_stops = p_formula.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)

        run = p_formula.add_run()
        run.add_tab()

        try:
            caminho_imagem_valido = formula_obj.caminho_processado_png
            if not caminho_imagem_valido or not os.path.exists(caminho_imagem_valido):
                    raise FileNotFoundError(f"Arquivo de imagem da fórmula não encontrado em '{caminho_imagem_valido}'")

            run.add_picture(caminho_imagem_valido, width=Cm(formula_obj.largura_cm))

        except Exception as e:
            erro_msg = f"[ERRO: Imagem da fórmula '{formula_obj.legenda}' não encontrada: {e}]"
            run_erro = p_formula.add_run(erro_msg)
            run_erro.italic = True

        run.add_tab()

        run_numero = p_formula.add_run(f"({formula_obj.numero})")
        self.regras._aplicar_formatacao_run(run_numero)

        p_legenda = self.doc.add_paragraph()
        p_legenda.add_run(f"Equação {formula_obj.numero} – {formula_obj.legenda}")
        self.regras.aplicar_estilo_legenda(p_legenda, is_titulo=True)
        p_legenda.paragraph_format.space_before = Pt(6)
        p_legenda.paragraph_format.space_after = Pt(12)

    def _get_marcador_lista(self, tipo: str, nivel: int, indice: int, prefixo_num: str) -> tuple[str, str]:
        """
        Retorna o marcador (ex: "a)") e o prefixo numérico para o próximo nível.
        """
        if tipo == "Híbrida (ABNT)":
            if nivel == 1: return f"{self._get_char_alfabetico(indice)})", "" # a)
            if nivel == 2: return f"{indice + 1})", ""                       # 1)
            if nivel == 3: return f"{self._get_char_romano(indice)}", ""     # i)
            return "-", ""                                                   # -
        
        if tipo == "Numérica (Seção)":
            num_atual = f"{prefixo_num}{indice + 1}."
            return num_atual, num_atual # Retorna o num_atual como prefixo para os filhos
            
        if tipo == "Alfabética":
            return f"{self._get_char_alfabetico(indice, maiusculo=True)}.", "" # A.
            
        if tipo == "Símbolos":
            return "•", "" # • (ou "-")
            
        return "", ""

    def _get_char_alfabetico(self, indice: int, maiusculo: bool = False) -> str:
        """Retorna 'a' para 0, 'b' para 1, etc."""
        offset = 65 if maiusculo else 97
        return chr(offset + indice)
        
    def _get_char_romano(self, indice: int) -> str:
        """Retorna 'i' para 0, 'ii' para 1, etc. (simplificado)"""
        romanos = ['i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x']
        return romanos[indice] if indice < len(romanos) else str(indice + 1)


    def _renderizar_lista(self, lista_obj: ListaABNT):
        """Renderiza a lista hierárquica no DOCX."""
        
        if lista_obj.mostrar_titulo:
            p_titulo = self.doc.add_paragraph()
            p_titulo.add_run(f"Lista {lista_obj.numero} – {lista_obj.titulo}")
            self.regras.aplicar_estilo_legenda(p_titulo, is_titulo=True)
            p_titulo.paragraph_format.keep_with_next = True

        def render_itens_recursivo(item_pai: ItemLista, nivel: int, prefixo_num: str):
            recuo_base_cm = 1.25
            recuo_nivel_cm = 0.75
            
            recuo_total_cm = recuo_base_cm + (nivel * recuo_nivel_cm)
            recuo_marcador_cm = recuo_base_cm + ((nivel - 1) * recuo_nivel_cm)

            if lista_obj.tipo_enumeracao == "Numérica (Seção)":
                recuo_total_cm = 1.25 
                recuo_marcador_cm = 0

            for i, item_filho in enumerate(item_pai.filhos):
                marcador, proximo_prefixo = self._get_marcador_lista(
                    lista_obj.tipo_enumeracao, nivel, i, prefixo_num
                )
                
                p = self.doc.add_paragraph()
                p.style = 'Normal' 
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
                p.paragraph_format.left_indent = Cm(recuo_total_cm)
                
                p.paragraph_format.first_line_indent = -Cm(recuo_total_cm - recuo_marcador_cm)
                
                if lista_obj.tipo_enumeracao == "Numérica (Seção)":
                    run = p.add_run(f"{marcador} {item_filho.texto}")
                    self.regras._aplicar_formatacao_run(run)
                else:
                    run_marcador = p.add_run(f"{marcador}\t")
                    self.regras._aplicar_formatacao_run(run_marcador)
                    
                    run_texto = p.add_run(item_filho.texto)
                    self.regras._aplicar_formatacao_run(run_texto)
                
                render_itens_recursivo(item_filho, nivel + 1, proximo_prefixo)

        render_itens_recursivo(lista_obj.raiz, 1, "")
        
    def _set_page_numbering(self, section):
        section.header.is_linked_to_previous = False
        header_p = section.header.paragraphs[0]
        header_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = header_p.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._r.append(instrText)
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_end)

    def _renderizar_capa(self):
        cfg = self.doc_abnt.configuracoes
        
        if cfg.posicao_brasao == "Lados (Esquerdo e Direito)":
            table_header = self.doc.add_table(rows=1, cols=3)
            table_header.columns[0].width = Cm(3.5); table_header.columns[1].width = Cm(9.0); table_header.columns[2].width = Cm(3.5)
            self._set_no_border_to_table(table_header)
            cell_L, cell_C, cell_R = table_header.cell(0,0), table_header.cell(0,1), table_header.cell(0,2)
            if cfg.caminho_brasao_esquerdo_processado and os.path.exists(cfg.caminho_brasao_esquerdo_processado):
                p_l = cell_L.paragraphs[0]; p_l.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p_l.add_run().add_picture(cfg.caminho_brasao_esquerdo_processado, width=Cm(cfg.tamanho_brasao_esquerdo_cm))
            p_c = cell_C.paragraphs[0]; p_c.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p_c.add_run(cfg.instituicao.upper()).bold = True
            if cfg.caminho_brasao_direito_processado and os.path.exists(cfg.caminho_brasao_direito_processado):
                p_r = cell_R.paragraphs[0]; p_r.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                p_r.add_run().add_picture(cfg.caminho_brasao_direito_processado, width=Cm(cfg.tamanho_brasao_direito_cm))
        else:
            p_inst = self.doc.add_paragraph() 
            p_inst.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if cfg.posicao_brasao == "Acima do Nome" and cfg.caminho_brasao_esquerdo_processado and os.path.exists(cfg.caminho_brasao_esquerdo_processado):
                p_inst.add_run().add_picture(cfg.caminho_brasao_esquerdo_processado, width=Cm(cfg.tamanho_brasao_esquerdo_cm))
                p_inst.add_run('\n')
            p_inst.add_run(cfg.instituicao.upper()).bold = True

        nomes_autores = '\n'.join([a.nome_completo.upper() for a in self.doc_abnt.autores])
        p_autores = self.doc.add_paragraph() 
        p_autores.paragraph_format.space_before = Pt(72)
        p_autores.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_autores.add_run(nomes_autores).bold = True

        p_titulo = self.doc.add_paragraph() 
        p_titulo.paragraph_format.space_before = Pt(120)
        p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_titulo = p_titulo.add_run(self.doc_abnt.titulo.upper())
        run_titulo.bold = True
        run_titulo.font.size = self.regras.TAMANHO_FONTE_CAPA
        
        section = self.doc.sections[0]
        footer = section.footer
        p_final = footer.paragraphs[0]
        p_final.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_final.add_run(f"{cfg.cidade.upper()}\n{cfg.ano}")


    def _renderizar_folha_rosto(self, section):
        cfg = self.doc_abnt.configuracoes
        
        nomes_autores = '\n'.join([a.nome_completo.upper() for a in self.doc_abnt.autores])
        p_autores = self.doc.add_paragraph()
        p_autores.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_autores.add_run(nomes_autores).bold = True

        p_titulo = self.doc.add_paragraph()
        p_titulo.paragraph_format.space_before = Pt(80)
        p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_titulo = p_titulo.add_run(self.doc_abnt.titulo.upper())
        run_titulo.bold = True
        run_titulo.font.size = self.regras.TAMANHO_FONTE_CAPA

        p_natureza = self.doc.add_paragraph()
        p_natureza.paragraph_format.space_before = Pt(90)
        texto_natureza = (f"{cfg.tipo_trabalho} apresentado ao curso de {cfg.modalidade_curso} em {cfg.curso} da {cfg.instituicao}, "
                            f"como requisito parcial para a obtenção do título de {cfg.titulo_pretendido}.")
        self.regras.aplicar_estilo_natureza_trabalho(p_natureza, texto_natureza)
        
        p_orientador = self.doc.add_paragraph()
        p_orientador.paragraph_format.space_before = Pt(12)
        self.regras.aplicar_estilo_natureza_trabalho(p_orientador, f"Orientador(a): {self.doc_abnt.orientador}")

        footer = section.footer
        p_final = footer.paragraphs[0]
        p_final.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_final.add_run(f"{cfg.cidade.upper()}\n{cfg.ano}")

    
    def _renderizar_resumo(self):
        self.regras.aplicar_estilo_titulo_secao(self.doc, numero="", titulo_texto="RESUMO")
        
        p_resumo = self.doc.add_paragraph()
        self.regras.aplicar_estilo_resumo(p_resumo, self.doc_abnt.resumo)
        
        p_resumo.paragraph_format.first_line_indent = self.regras.RECUO_PRIMEIRA_LINHA

        self.doc.add_paragraph() 
        
        p_kw = self.doc.add_paragraph()
        run_kw = p_kw.add_run("Palavras-chave: ")
        run_kw.bold = True
        
        texto_kw = self.doc_abnt.palavras_chave.replace(';', '.') + "."
        p_kw.add_run(texto_kw)

    def _renderizar_sumario(self):
        self.regras.aplicar_estilo_titulo_secao(self.doc, numero="", titulo_texto="SUMÁRIO")
        paragrafo_placeholder = self.doc.add_paragraph()
        adicionar_sumario(self.doc, paragrafo_placeholder)
        self.doc.add_page_break()

    def _renderizar_referencias(self):
        self.regras.aplicar_estilo_titulo_secao(self.doc, numero="", titulo_texto="REFERÊNCIAS")
        self.doc_abnt.ordenar_referencias()
        for ref in self.doc_abnt.referencias:
            p_ref = self.doc.add_paragraph()
            self.regras.aplicar_estilo_referencia(p_ref, ref.formatar())